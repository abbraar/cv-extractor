"""Extract plain text from CV files (PDF, DOCX, TXT). PDF: text layer + optional OCR for scans."""

from __future__ import annotations

import io
import os
from pathlib import Path

import pdfplumber
from docx import Document

# Filled when Gemini vision runs on a PDF; read via get_last_pdf_vision_diagnostic() for UI.
_last_pdf_vision_diagnostic: list[str] = []


def get_last_pdf_vision_diagnostic() -> str:
    return "\n".join(_last_pdf_vision_diagnostic).strip()


def _extract_pdf_text_layer(data: bytes) -> str:
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
    return "\n\n".join(parts).strip()


def _windows_tesseract_exe() -> str | None:
    if os.name != "nt":
        return None
    for path in (
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ):
        if Path(path).is_file():
            return path
    return None


def _configure_tesseract() -> None:
    try:
        import pytesseract
    except ImportError:
        return
    cmd = os.environ.get("TESSERACT_CMD", "").strip()
    if not cmd:
        cmd = _windows_tesseract_exe() or ""
    if cmd:
        pytesseract.pytesseract.tesseract_cmd = cmd


def _extract_pdf_ocr(data: bytes) -> str:
    """
    OCR for image-only / scanned PDFs using PyMuPDF (render) + Tesseract.
    Requires: pip install pymupdf pytesseract Pillow, and Tesseract installed on the system.
    """
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""

    _configure_tesseract()

    try:
        pytesseract.get_tesseract_version()
    except Exception:
        return ""

    parts: list[str] = []
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        for page in doc:
            # Higher DPI improves OCR on small text; cap memory on huge pages
            mat = fitz.Matrix(2.0, 2.0)  # ~144 DPI base → ~288 effective scale
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_bytes))

            text = ""
            for lang in ("eng+ara", "ara+eng", "eng", "ara"):
                try:
                    text = pytesseract.image_to_string(img, lang=lang)
                except Exception:
                    continue
                if text and text.strip():
                    break
            if text.strip():
                parts.append(text.strip())
    finally:
        doc.close()

    return "\n\n".join(parts).strip()


def _extract_pdf(data: bytes) -> str:
    text = _extract_pdf_text_layer(data)
    if text:
        return text
    text = _extract_pdf_ocr(data)
    if text:
        return text
    # Scanned PDFs: OpenAI vision if OPENAI_API_KEY is set.
    from openai_service import extract_text_from_scanned_pdf_via_openai

    _last_pdf_vision_diagnostic.clear()
    return extract_text_from_scanned_pdf_via_openai(
        data, diagnostic_lines=_last_pdf_vision_diagnostic
    )


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Route by extension and return UTF-8 text."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix == ".docx":
        return _extract_docx(data)
    if suffix == ".txt":
        return _extract_txt(data)
    raise ValueError(f"Unsupported file type: {suffix}. Use .pdf, .docx, or .txt.")


def pdf_ocr_dependencies_ok() -> bool:
    """True if Tesseract OCR path can run (for UI hints)."""
    try:
        import pytesseract  # noqa: F401
        from PIL import Image  # noqa: F401
    except ImportError:
        return False
    _configure_tesseract()
    try:
        import pytesseract as pt

        pt.get_tesseract_version()
        return True
    except Exception:
        return False


def pdf_openai_vision_ready() -> bool:
    """True if scanned PDFs can use OpenAI vision (PyMuPDF + API key)."""
    if not (os.environ.get("OPENAI_API_KEY") or "").strip():
        return False
    try:
        import fitz  # noqa: F401
    except ImportError:
        return False
    return True


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    # Many CV templates put content only in tables; paragraphs stay empty.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts).strip()


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1256", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace").strip()
